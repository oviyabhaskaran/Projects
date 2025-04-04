import streamlit as st
import requests
from config import NEWSAPI_API_KEY
from transformers import pipeline
import xml.etree.ElementTree as ET
from docx import Document
from io import BytesIO

# Initialize the summarizer
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# Streamlit App UI
st.title("InsightSnap: Instant News & Research Summary Hub")
st.markdown("---")

# Search Input
query = st.text_input("Enter your search query:")
if st.button("Generate"):
    if query:
        with st.spinner("Fetching and summarizing..."):
            
            def fetch_news(query, num_articles=3):
                url = f"https://newsapi.org/v2/everything?q={query}&pageSize={num_articles}&apiKey={NEWSAPI_API_KEY}"
                response = requests.get(url)
                return response.json().get("articles", []) if response.status_code == 200 else []
            
            def summarize_article(article_text):
                return summarizer(article_text, max_length=150, min_length=50, do_sample=False)[0]['summary_text'] if article_text else "No content available."
            
            def fetch_research_papers(query, max_results=3):
                base_url = "http://export.arxiv.org/api/query"
                params = {"search_query": query, "start": 0, "max_results": max_results}
                response = requests.get(base_url, params=params)
                
                if response.status_code == 200:
                    root = ET.fromstring(response.text)
                    papers = []
                    for entry in root.findall("{http://www.w3.org/2005/Atom}entry"):
                        papers.append({
                            "title": entry.find("{http://www.w3.org/2005/Atom}title").text,
                            "summary": entry.find("{http://www.w3.org/2005/Atom}summary").text,
                            "url": entry.find("{http://www.w3.org/2005/Atom}id").text,
                        })
                    return papers
                return []
            
            def generate_docx(news, papers):
                doc = Document()
                doc.add_heading("InsightSnap: Summary Report", 0)
                doc.add_heading("News Summaries", level=1)
                
                for article in news:
                    doc.add_paragraph(f"**{article['title']}**")
                    doc.add_paragraph(f"{article['summary']}")
                    doc.add_paragraph("---")
                
                doc.add_heading("Research Paper Summaries", level=1)
                
                for paper in papers:
                    doc.add_paragraph(f"**{paper['title']}**")
                    doc.add_paragraph(f"{paper['summary']}")
                    doc.add_paragraph(f"URL: {paper['url']}")
                    doc.add_paragraph("---")
                
                # Save the document in memory
                docx_io = BytesIO()
                doc.save(docx_io)
                docx_io.seek(0)  
                return docx_io
            
            # Fetch and summarize news and research papers
            news_articles = fetch_news(query)
            research_papers = fetch_research_papers(query)
            
            summarized_news = [{"title": art.get("title", "No title"), "summary": summarize_article(art.get("content", ""))} for art in news_articles]
            summarized_papers = [{"title": p.get("title", "No title"), "summary": summarize_article(p.get("summary", "")), "url": p.get("url", "No URL")} for p in research_papers]
            
            # Display the summaries in Streamlit
            st.subheader("News Summaries")
            for article in summarized_news:
                st.write(f"**{article['title']}**")
                st.write(article['summary'])
                st.markdown("---")
            
            st.subheader("Research Paper Summaries")
            for paper in summarized_papers:
                st.write(f"**{paper['title']}**")
                st.write(paper['summary'])
                st.write(f"**URL:** {paper['url']}")
                st.markdown("---")
            
            # Generate docx file in memory
            docx_data = generate_docx(summarized_news, summarized_papers)
            
            st.download_button("Download Report", data=docx_data, file_name="summary_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Please enter a query to generate summaries.")
