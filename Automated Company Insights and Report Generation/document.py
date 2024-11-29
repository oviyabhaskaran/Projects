import os
import json
from docx import Document
from docx.shared import Inches


def create_company_document(data_dict, document, image_path=None):
    # Assuming 'data_dict' structure is correctly defined before calling this function
    # Add sections dynamically
    for section_title, content in data_dict.items():
        if section_title == "Company Name":  # Handle Company Name differently
            # print('Company Name:',content)
            document.add_heading(content, level=1)
            # for value in content:
            #     document.add_heading(value, level=1)
        elif section_title == "Overall Summary":  # Handle Overall Summary if special treatment needed
            document.add_paragraph(content)
            # for value in content:
            #     document.add_paragraph(value)
        else:  # Handle all other sections with bullets
            document.add_heading(f'{section_title}:', level=2)  # Adjusted for consistency, making titles level 2
            if isinstance(content, list):  # Checking if content is list to iterate
                for value in content:
                    document.add_paragraph(value, style='List Bullet')
            else:  # If content is not a list (single item), just add it
                document.add_paragraph(content, style='List Bullet')

    # Add an image if 'image_path' is provided and is not None
    if image_path:
        document.add_picture(image_path, width=Inches(6.5), height=Inches(3.5))


def read_json(file_path):
    """
    Read data from a JSON file and return it as a Python dictionary.

    Parameters:
        file_path (str): The path to the JSON file.

    Returns:
        dict: The data read from the JSON file.
    """
    # Read the JSON file and load its contents into a dictionary
    with open(file_path, "r") as json_file:
        data_dict = json.load(json_file)

    return data_dict

def generate_company_report_document(json_directory, output_plot_directories):
    """
    Generate a company report document based on JSON data and associated plot images.

    Parameters:
        json_directory (str): Path to the directory containing JSON files.
        output_plot_directories (list of str): List of paths to directories containing plot images.
    """
    # Create a new document
    document = Document()

    # Iterate through each file in the JSON directory
    for filename in os.listdir(json_directory):
        # Check if the file is a JSON file
        if filename.endswith(".json"):
            # Construct the full file path
            file_path = os.path.join(json_directory, filename)

            # Read the data from the JSON file
            data_dict = read_json(file_path)

            # Extract the image path from the output plot directory 1
            image_filename1 = filename.split('.')[0] + ".png"
            image_path1 = os.path.join(output_plot_directories[0], image_filename1)

            # Create a section for the current JSON file and add output_plot_directory1 image
            create_company_document(data_dict, document, image_path1)

            # Extract the image path from the output plot directory 2
            image_path2 = os.path.join(output_plot_directories[1], image_filename1)

            # Add output_plot_directory2 image after output_plot_directory1 image
            if os.path.exists(image_path2):
                document.add_picture(image_path2, width=Inches(6.5), height=Inches(3.5))

    # Save the document inside a directory
    directory_name = "output_company_reports"
    if not os.path.exists(directory_name):
        os.makedirs(directory_name)

    document_name = os.path.join(directory_name, "Company_Performance_Report.docx")
    document.save(document_name)
    return directory_name
